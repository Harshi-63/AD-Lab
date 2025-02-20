import os
import logging
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from llama_index.core import VectorStoreIndex
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama
from llama_index.core.schema import Document
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.settings import Settings

# Initialize Flask app
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# Set max request size (10MB for file uploads)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024  

# Define upload folder
UPLOAD_FOLDER = "uploads"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Logging setup
logging.basicConfig(level=logging.DEBUG)

# Load Local Embedding Model
embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Initialize Ollama as Local LLM
try:
    llm = Ollama(model="mistral", request_timeout=300, temperature=0.7, num_ctx=512)  # Change model if needed
    logging.info("Ollama model loaded successfully.")
except Exception as e:
    logging.error(f"Error loading LLM: {str(e)}")
    llm = None  # Prevent crashes

# Configure global settings
Settings.embed_model = embed_model
Settings.llm = llm

# Store documents globally
documents = []
index = None


def extract_text_from_file(file_path):
    """Extract text from PDF, Word, or Excel files."""
    from PyPDF2 import PdfReader
    import docx
    import pandas as pd

    text = ""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif ext == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif ext in [".xls", ".xlsx"]:
            df = pd.read_excel(file_path)
            text = df.to_string()
        else:
            text = "Unsupported file format."
    except Exception as e:
        text = f"Error extracting text: {str(e)}"
        logging.error(f"File extraction error: {text}")

    return text


@app.route("/", methods=["GET"])
def home():
    return jsonify({"message": "Flask Backend is Running!"})


@app.route("/upload", methods=["POST"])
def upload_file():
    global index, documents

    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    # Secure filename and save it
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(file_path)

    logging.info(f"File uploaded: {filename}")

    # Extract text from the document
    extracted_text = extract_text_from_file(file_path)

    if "Error extracting text" in extracted_text:
        return jsonify({"error": extracted_text}), 500

    # Store extracted text as a document
    document = Document(text=extracted_text)
    documents.append(document)

    # Create or update the index
    index = VectorStoreIndex.from_documents(documents)

    return jsonify({"message": "File uploaded and processed successfully!"})


@app.route("/chat", methods=["POST"])
def chat():
    global index

    if index is None:
        return jsonify({"error": "No documents have been uploaded yet"}), 400

    data = request.get_json()
    query_text = data.get("query")

    if not query_text:
        return jsonify({"error": "Query text is required"}), 400

    try:
        retriever = index.as_retriever()
        query_engine = RetrieverQueryEngine(retriever=retriever)

        logging.debug(f"Processing query: {query_text}")

        response = query_engine.query(query_text)

        logging.debug(f"Response: {response}")

        return jsonify({"response": str(response)})

    except Exception as e:
        logging.error(f"Error processing query: {str(e)}")
        return jsonify({"error": f"Error processing query: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
